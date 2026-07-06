"""Generate synthetic test fixtures. Run once: python tests/generate_fixtures.py"""
import json
from pathlib import Path

import fitz
from docx import Document

FIXTURES = Path(__file__).parent / "fixtures"


def make_caste_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    text = (
        "CASTE CERTIFICATE\n\n"
        "Government of Maharashtra\n"
        "This is to certify that the applicant belongs to Scheduled Caste category.\n\n"
        "Name: Rajesh Kumar\n"
        "Caste: Scheduled Caste\n"
        "Certificate No: SC/2024/12345\n"
        "Issue Date: 15-03-2024\n"
        "Issuing Authority: Tahsildar Office, Pune\n"
    )
    page.insert_text((72, 72), text, fontsize=11)
    doc.save(str(path))
    doc.close()


def make_experience_docx(path: Path, include_email: bool = True) -> None:
    doc = Document()
    doc.add_heading("Experience Certificate", level=1)
    doc.add_paragraph(
        "Acme Technologies Pvt. Ltd. hereby certify that Mr. John Doe was employed "
        "with our organization as Senior Engineer from 01-01-2020 to 31-12-2023."
    )
    if include_email:
        doc.add_paragraph("For verification contact: hr@acme.com")
    doc.add_paragraph("Designation: Senior Engineer")
    doc.save(str(path))


def make_manifest(path: Path) -> None:
    manifest = {
        "caste_sample.pdf": {
            "doc_type": "caste",
            "min_score": 50,
            "required_flags_absent": ["TEXT_EXTRACT_FAILED"],
            "required_fields": ["caste_category"],
        },
        "experience_sample.docx": {
            "doc_type": "experience",
            "min_score": 40,
            "required_fields": ["hr_email", "company_name"],
            "required_flags_absent": ["NO_EMAIL"],
        },
        "experience_no_email.docx": {
            "doc_type": "experience",
            "required_flags": ["NO_EMAIL"],
            "expected_status": "Red Flagged",
        },
    }
    path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def main() -> None:
    FIXTURES.mkdir(parents=True, exist_ok=True)
    (FIXTURES / "real").mkdir(exist_ok=True)
    make_caste_pdf(FIXTURES / "caste_sample.pdf")
    make_experience_docx(FIXTURES / "experience_sample.docx", include_email=True)
    make_experience_docx(FIXTURES / "experience_no_email.docx", include_email=False)
    make_manifest(FIXTURES / "manifest.json")
    print(f"Fixtures written to {FIXTURES}")


if __name__ == "__main__":
    main()
